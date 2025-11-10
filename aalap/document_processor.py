#!/usr/bin/env python3
"""
Document processing utilities for Aalap RAG
Handles various document formats for indexing
"""

from pathlib import Path
from typing import Dict, Optional, List
import json

class DocumentProcessor:
    """Process various document types for RAG indexing"""

    @staticmethod
    def process_text_file(file_path: Path) -> tuple[str, Dict]:
        """Process plain text files"""
        content = file_path.read_text(encoding='utf-8', errors='ignore')
        metadata = {
            "source": str(file_path),
            "filename": file_path.name,
            "type": "text",
            "size": file_path.stat().st_size
        }
        return content, metadata

    @staticmethod
    def process_pdf(file_path: Path) -> tuple[str, Dict]:
        """Process PDF files"""
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(str(file_path))
            content = ""
            for page in reader.pages:
                content += page.extract_text() + "\n"

            metadata = {
                "source": str(file_path),
                "filename": file_path.name,
                "type": "pdf",
                "pages": len(reader.pages)
            }
            return content, metadata
        except ImportError:
            raise ImportError("PyPDF2 not installed. Install with: pip install PyPDF2")

    @staticmethod
    def process_docx(file_path: Path) -> tuple[str, Dict]:
        """Process Word documents"""
        try:
            from docx import Document
            doc = Document(str(file_path))
            content = "\n".join([paragraph.text for paragraph in doc.paragraphs])

            metadata = {
                "source": str(file_path),
                "filename": file_path.name,
                "type": "docx",
                "paragraphs": len(doc.paragraphs)
            }
            return content, metadata
        except ImportError:
            raise ImportError("python-docx not installed. Install with: pip install python-docx")

    @staticmethod
    def process_markdown(file_path: Path) -> tuple[str, Dict]:
        """Process Markdown files"""
        content = file_path.read_text(encoding='utf-8')

        # Extract title if present
        title = None
        lines = content.split('\n')
        if lines and lines[0].startswith('#'):
            title = lines[0].lstrip('#').strip()

        metadata = {
            "source": str(file_path),
            "filename": file_path.name,
            "type": "markdown",
            "title": title
        }
        return content, metadata

    @staticmethod
    def process_code(file_path: Path) -> tuple[str, Dict]:
        """Process source code files"""
        content = file_path.read_text(encoding='utf-8', errors='ignore')

        # Detect language from extension
        extension_to_lang = {
            '.py': 'python', '.js': 'javascript', '.ts': 'typescript',
            '.java': 'java', '.cpp': 'cpp', '.c': 'c', '.h': 'c',
            '.cs': 'csharp', '.go': 'go', '.rs': 'rust', '.rb': 'ruby',
            '.php': 'php', '.swift': 'swift', '.kt': 'kotlin'
        }

        language = extension_to_lang.get(file_path.suffix.lower(), 'unknown')

        metadata = {
            "source": str(file_path),
            "filename": file_path.name,
            "type": "code",
            "language": language,
            "lines": len(content.split('\n'))
        }
        return content, metadata

    @staticmethod
    def process_json(file_path: Path) -> tuple[str, Dict]:
        """Process JSON files"""
        with open(file_path, 'r') as f:
            data = json.load(f)

        # Convert to readable text
        content = json.dumps(data, indent=2)

        metadata = {
            "source": str(file_path),
            "filename": file_path.name,
            "type": "json"
        }
        return content, metadata

    @classmethod
    def process_file(cls, file_path: Path) -> Optional[tuple[str, Dict]]:
        """
        Process a file and return content and metadata

        Args:
            file_path: Path to file

        Returns:
            (content, metadata) tuple or None if unsupported
        """
        suffix = file_path.suffix.lower()

        processors = {
            '.txt': cls.process_text_file,
            '.md': cls.process_markdown,
            '.pdf': cls.process_pdf,
            '.docx': cls.process_docx,
            '.json': cls.process_json,
            '.py': cls.process_code,
            '.js': cls.process_code,
            '.ts': cls.process_code,
            '.java': cls.process_code,
            '.cpp': cls.process_code,
            '.c': cls.process_code,
            '.h': cls.process_code,
            '.cs': cls.process_code,
            '.go': cls.process_code,
            '.rs': cls.process_code,
            '.rb': cls.process_code,
            '.php': cls.process_code,
            '.swift': cls.process_code,
            '.kt': cls.process_code,
        }

        processor = processors.get(suffix)
        if processor:
            try:
                return processor(file_path)
            except Exception as e:
                print(f"Error processing {file_path}: {e}")
                return None

        return None

    @classmethod
    def process_directory(
            cls,
            dir_path: Path,
            recursive: bool = True,
            extensions: Optional[List[str]] = None
    ) -> List[tuple[Path, str, Dict]]:
        """
        Process all files in a directory

        Args:
            dir_path: Directory path
            recursive: Whether to process subdirectories
            extensions: List of extensions to process (None = all supported)

        Returns:
            List of (file_path, content, metadata) tuples
        """
        results = []

        pattern = "**/*" if recursive else "*"
        for file_path in dir_path.glob(pattern):
            if not file_path.is_file():
                continue

            if extensions and file_path.suffix.lower() not in extensions:
                continue

            result = cls.process_file(file_path)
            if result:
                content, metadata = result
                results.append((file_path, content, metadata))

        return results